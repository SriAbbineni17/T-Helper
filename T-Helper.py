from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor
import requests
import json
import tkinter as tk
from tkinter import ttk
from tkinter.filedialog import askopenfilename, asksaveasfilename
from docx.shared import Pt
import os

def open_file():
    """Open a file for editing."""

    filepath = askopenfilename(
        filetypes=[("Docx Files", "*.docx"), ("All Files", "*.*")]
    )
    if not filepath:
        return
    filename = filepath[filepath.rfind('/')+1:len(filepath)]
    file_queue.append(filename)
    label = tk.Label(fr_queue, text=filename).pack(side='top', fill = 'both')
    progress = tk.ttk.Progressbar(fr_queue, orient = 'horizontal', length = 100, mode = 'determinate')
    progress.pack(side='top', fill = 'both') 
    # print(filename)
    url = "https://deep-translate1.p.rapidapi.com/language/translate/v2"

    querystring = {"to[0]":"te","api-version":"3.0","profanityAction":"NoAction","textType":"plain"}

    doc = Document(filepath)
    table = doc.tables[0]
    # len(table.rows)
    
    for i in range(4, len(table.rows)):
        line = table.rows[i].cells[1].text
#sri key: 926576883emsh0929c234d91e5d8p19a217jsn1b07e87e770d

        # print(line)
        payload = {
            "q": line,
            "source": "en",
	        "target": "te"        
        }
        headers = {
            "content-type": "application/json",
            "X-RapidAPI-Key": "6fbae1f9d6mshb826f520c2f1e36p1b376cjsnf532747f5544",
            "X-RapidAPI-Host": "deep-translate1.p.rapidapi.com"
        }
        response = requests.post(url, json=payload, headers=headers)
        # print(i)
        
        r = response.json()
        # print(r)
        table.rows[i].cells[2].text = r['data']['translations']['translatedText']
        # print(r['data']['translations']['translatedText'])
        # run = table.rows[i].cells[2].add_paragraph().add_run(r[0]["translations"][0]["text"])
        run = table.rows[i].cells[2].paragraphs[0].runs[0]
        run.font.name = 'Mallanna'
        run.font.size = Pt(12)
        progress['value'] = i/len(table.rows) * progress['length']
        window.update_idletasks()
    
    
    doc.save(filepath)
    new_name = filepath[0:-5] + "Telugu" + ".docx"
    # print(new_name)
    os.rename(filepath, new_name)
    label = tk.Label(fr_queue, text="Done, open your file!").pack(side='top', fill = 'both')
file_queue = []
window = tk.Tk()
window.title("T-Helper")
window.rowconfigure(0, minsize=900, weight=1)
window.columnconfigure(1, minsize=900, weight=1)
fr_buttons = tk.Frame(window, relief=tk.RIDGE, bd=2)
fr_queue = tk.Frame(window, relief=tk.RIDGE, bd=2)
btn_open = tk.Button(fr_buttons, text="Open", command=open_file)

fr_buttons.pack(side='left', fill = 'both')
# .grid(row=0, column=0, sticky="ns")
fr_queue.pack(side='left', fill = 'both', expand=True)
# .grid(row=0, column=1, sticky="nw")
btn_open.pack(side='left', fill = 'both')
# .grid(row=0, column=0, sticky="ew", padx=5, pady=5)

window.mainloop()



